import pandas as pd
from docx import Document # python-docx
from docx.shared import Pt
from docx.oxml.ns import qn

# 读取CSV数据
input_path = "../B_novel_analysis/output/test_结构化数据_不含情绪动作.csv"
df = pd.read_csv(input_path)

# 创建Word文档
doc = Document()
doc.styles['Normal'].font.name = u'宋体'  # 设置中文字体
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

# 遍历每一行，添加到文档中
for _, row in df.iterrows():
    role = str(row["role"]).strip()
    dialogue = str(row["emo_label"]).strip()

    dialogue = dialogue.replace(r'\n\n', '')
    #print(f"处理角色：{role}，对白：{dialogue}")

    # 跳过空dialogue
    if not dialogue or dialogue.startswith("(生成失败"):
        continue

    if role in ["", "旁白"]:
        para = doc.add_paragraph()
        run = para.add_run(dialogue)
        run.font.size = Pt(10)
        run.bold = False
    else:
        para = doc.add_paragraph()
        run = para.add_run(f"{role}：{dialogue}")
        run.font.size = Pt(12)
        run.bold = True

# 保存文档
output_path = "剧本输出_龙傲天版.docx"
doc.save(output_path)
print(f"✅ Word剧本已保存到：{output_path}")
